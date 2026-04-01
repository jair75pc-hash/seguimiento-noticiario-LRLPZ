import streamlit as st
from docx import Document
from docx.shared import Inches
from streamlit_paste_button import paste_image_button
import io
import matplotlib.pyplot as plt

st.set_page_config(page_title="Monitor de Noticias - CECYTEM", layout="wide")
st.title("📊 Monitor de Noticias Inteligente (Servicio Social)")

# --- Memoria ---
if 'datos_mes' not in st.session_state: st.session_state.datos_mes = {}
if 'dia' not in st.session_state: st.session_state.dia = 1

# --- Navegación ---
c1, c2, c3 = st.columns([1, 2, 1])
with c1:
    if st.button("⬅️ Anterior"):
        if st.session_state.dia > 1: st.session_state.dia -= 1
with c2:
    st.markdown(f"<h1 style='text-align: center;'>📅 {st.session_state.dia} de Marzo</h1>", unsafe_allow_html=True)
with c3:
    if st.button("Siguiente ➡️"):
        if st.session_state.dia < 31: st.session_state.dia += 1

st.divider()

key_dia = str(st.session_state.dia)
if key_dia not in st.session_state.datos_mes:
    st.session_state.datos_mes[key_dia] = []

if st.button("➕ AGREGAR NOTICIA"):
    st.session_state.datos_mes[key_dia].append({"titulo": "", "img": None, "sent": "Neutro", "link": ""})

# --- Listado de Noticias ---
for i, noticia in enumerate(st.session_state.datos_mes[key_dia]):
    with st.container(border=True):
        col_t, col_b = st.columns([5, 1])
        with col_t: st.subheader(f"Noticia {i+1}")
        with col_b: 
            if st.button("🗑️", key=f"del_{key_dia}_{i}"):
                st.session_state.datos_mes[key_dia].pop(i)
                st.rerun()

        c_img, c_txt = st.columns([1, 1])
        with c_img:
            st.write("👉 **Paso:** Captura pantalla y pica abajo:")
            # BOTÓN ESPECIAL DE PEGADO
            pasted = paste_image_button(label="📋 PEGAR CAPTURA AQUÍ", key=f"p_{key_dia}_{i}")
            if pasted and pasted.image_data is not None:
                noticia["img"] = pasted.image_data
            
            if noticia["img"] is not None:
                st.image(noticia["img"], caption="Vista previa", width=300)
            
        with c_txt:
            noticia["titulo"] = st.text_input("📝 Título", value=noticia["titulo"], key=f"t_{key_dia}_{i}")
            noticia["sent"] = st.select_slider("😊 Sentimiento", options=["Negativo", "Neutro", "Positivo"], value=noticia["sent"], key=f"s_{key_dia}_{i}")
            noticia["link"] = st.text_input("🔗 Link", value=noticia["link"], key=f"l_{key_dia}_{i}")

st.divider()

# --- Generador de Word Profesional ---
def crear_reporte():
    doc = Document()
    doc.add_heading('Reporte de Seguimiento Mensual - Marzo 2026', 0)
    
    # Tabla de Índice
    doc.add_heading('Índice de Noticias', level=1)
    table = doc.add_table(rows=1, cols=3); table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'Día', 'Noticia', 'Sentimiento'
    
    stats = {"Positivo": 0, "Neutro": 0, "Negativo": 0}
    noticias_validas = []

    for d in range(1, 32):
        k = str(d)
        if k in st.session_state.datos_mes:
            for n in st.session_state.datos_mes[k]:
                row = table.add_row().cells
                row[0].text = f"{d}/Marzo"
                row[1].text = n["titulo"]
                row[2].text = n["sent"]
                stats[n["sent"]] += 1
                noticias_validas.append((d, n))

    doc.add_page_break()

    # Detalle con Imágenes
    for dia, n in noticias_validas:
        doc.add_heading(f"Día {dia}: {n['titulo']}", level=2)
        if n["img"] is not None:
            img_io = io.BytesIO()
            n["img"].save(img_io, format='PNG')
            doc.add_picture(img_io, width=Inches(5))
        doc.add_paragraph(f"Sentimiento: {n['sent']}\nLink: {n['link']}")
        doc.add_page_break()

    # Gráfica Final
    if sum(stats.values()) > 0:
        plt.figure(figsize=(5, 5))
        plt.pie(stats.values(), labels=stats.keys(), autopct='%1.1f%%', colors=['#ff9999','#66b3ff','#99ff99'])
        plt.title("Resumen de Sentimientos del Mes")
        img_plt = io.BytesIO()
        plt.savefig(img_plt, format='png')
        doc.add_heading('Gráfica de Impacto', level=1)
        doc.add_picture(img_plt, width=Inches(4.5))

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

if st.button("🚀 GENERAR REPORTE COMPLETO"):
    archivo = crear_reporte()
    st.download_button("📥 Descargar Reporte.docx", data=archivo, file_name="Reporte_Noticias_Final.docx")
    st.balloons()